#report_generator.py
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors
from reportlab.lib.units import inch
import io
import re
import os

# python-docx가 있으면 import, 없으면 None
try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    Document = None
    WD_ALIGN_PARAGRAPH = None

def register_korean_font():
    """한국어 폰트 등록 - 여러 옵션 시도"""
    font_options = [
        'NOTOSANSKR-VF.TTF',
        'NanumGothicCoding.ttf',
        'NanumGothicCoding-Bold.ttf',
        'malgun.ttf',  # Windows 기본 폰트
        'gulim.ttc',   # Windows 기본 폰트
    ]
    
    for font_file in font_options:
        try:
            if os.path.exists(font_file):
                pdfmetrics.registerFont(TTFont('KoreanFont', font_file))
                print(f"한국어 폰트 등록 성공: {font_file}")
                return True
        except Exception as e:
            print(f"폰트 등록 실패 ({font_file}): {e}")
            continue
    
    # 폰트 파일이 없으면 기본 폰트 사용
    print("한국어 폰트를 찾을 수 없어 기본 폰트를 사용합니다.")
    return False

def clean_text_for_pdf(text):
    """PDF용 텍스트 정리 - HTML 태그 제거 및 안전한 형식으로 변환 - 개선된 버전"""
    if not text:
        return ""
    
    # HTML 태그 제거
    text = re.sub(r'<br\s*/?>', '\n', text)  # <br> 태그를 줄바꿈으로
    text = re.sub(r'<[^>]+>', '', text)  # 모든 HTML 태그 제거
    
    # 특수 문자 처리
    text = text.replace('•', '•')  # bullet point
    text = text.replace('–', '-')  # en dash
    text = text.replace('—', '-')  # em dash
    text = text.replace('"', '"')  # smart quotes
    text = text.replace('"', '"')  # smart quotes
    text = text.replace(''', "'")  # smart apostrophe
    text = text.replace(''', "'")  # smart apostrophe
    
    # 표 관련 특수 문자 처리
    text = text.replace('│', '|')  # box drawing characters
    text = text.replace('┌', '')
    text = text.replace('┐', '')
    text = text.replace('└', '')
    text = text.replace('┘', '')
    text = text.replace('├', '')
    text = text.replace('┤', '')
    text = text.replace('┬', '')
    text = text.replace('┴', '')
    text = text.replace('─', '-')
    
    # 연속된 공백 정리 (표 셀 내에서는 보존)
    if '|' not in text:  # 표가 아닌 경우에만 공백 정리
        text = re.sub(r'\s+', ' ', text)
    
    # 줄바꿈 정리 (표가 아닌 경우에만)
    if '|' not in text:
        text = re.sub(r'\n\s*\n', '\n\n', text)
    
    return text.strip()

def parse_table_from_text(text):
    """텍스트에서 표 형식을 파싱하여 2D 배열로 변환 - 개선된 버전"""
    lines = text.strip().split('\n')
    table_data = []
    table_title = None
    
    # 표 제목 찾기 (표 위의 텍스트)
    title_lines = []
    table_started = False
    
    for i, line in enumerate(lines):
        line = line.strip()
        if not line:
            continue
            
        # 표 시작 확인
        if is_table_row(line):
            table_started = True
            # 이전까지의 텍스트를 제목으로 처리
            if title_lines:
                table_title = ' '.join(title_lines).strip()
                # 제목에서 불필요한 문자 제거
                table_title = re.sub(r'^\*\*|\*\*$', '', table_title)  # 마크다운 볼드 제거
                table_title = re.sub(r'^#+\s*', '', table_title)  # 마크다운 헤더 제거
            break
        else:
            # 표가 시작되기 전까지의 텍스트를 제목으로 저장
            # 단, 너무 긴 텍스트는 제목이 아닐 수 있음
            if len(line) < 100:  # 100자 이하만 제목으로 간주
                title_lines.append(line)
    
    # 표 데이터 파싱
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # 구분선 제거 (마크다운 표 구분선)
        if re.match(r'^[\s\-=_:|]+\s*$', line):
            continue
            
        # 표 행인지 확인
        if is_table_row(line):
            cells = parse_table_row(line)
            if cells:
                table_data.append(cells)
    
    # 표 데이터 정규화 (모든 행이 같은 열 수를 가지도록)
    if table_data:
        max_cols = max(len(row) for row in table_data)
        normalized_data = []
        for row in table_data:
            # 부족한 열은 빈 문자열로 채움
            normalized_row = row + [''] * (max_cols - len(row))
            normalized_data.append(normalized_row)
        return normalized_data, table_title
    
    return table_data, table_title

def is_table_row(line):
    """한 줄이 표 행인지 확인"""
    # | 구분자가 있는 경우
    if '|' in line:
        return True
    
    # 탭으로 구분된 경우
    if '\t' in line:
        return True
    
    # 2개 이상의 공백으로 구분된 경우 (정렬된 텍스트)
    if re.search(r'\s{2,}', line):
        return True
    
    return False

def parse_table_row(line):
    """표 행을 파싱하여 셀 배열로 변환"""
    # | 구분자로 분할 (마크다운 표 형식)
    if '|' in line:
        cells = [cell.strip() for cell in line.split('|')]
        # 첫 번째와 마지막 빈 셀 제거 (마크다운 표 형식)
        if cells and not cells[0].strip():
            cells = cells[1:]
        if cells and not cells[-1].strip():
            cells = cells[:-1]
        return cells
    
    # 탭으로 구분된 경우
    elif '\t' in line:
        cells = [cell.strip() for cell in line.split('\t') if cell.strip()]
        return cells
    
    # 2개 이상의 공백으로 구분된 경우
    elif re.search(r'\s{2,}', line):
        cells = [cell.strip() for cell in re.split(r'\s{2,}', line) if cell.strip()]
        return cells
    
    return []

def is_table_format(text):
    """텍스트가 표 형식인지 확인 - 개선된 버전"""
    lines = text.strip().split('\n')
    if len(lines) < 2:
        return False
    
    # 표 구분자 확인
    table_indicators = ['|', '\t']
    table_line_count = 0
    
    for line in lines[:5]:  # 처음 5줄만 확인
        if any(indicator in line for indicator in table_indicators):
            table_line_count += 1
    
    # 2줄 이상에 표 구분자가 있으면 표로 인식
    if table_line_count >= 2:
        return True
    
    # 구분선 확인 (마크다운 표 구분선)
    for line in lines:
        if re.match(r'^[\s\-=_:|]+\s*$', line.strip()):
            return True
    
    # 정렬된 텍스트 확인 (2개 이상의 공백으로 구분)
    aligned_line_count = 0
    for line in lines[:3]:
        if re.search(r'\s{2,}', line.strip()):
            aligned_line_count += 1
    
    if aligned_line_count >= 2:
        return True
    
    return False

def create_table_with_improved_style(table_data, font_registered):
    """개선된 스타일로 표 생성"""
    if not table_data or len(table_data) == 0:
        return None
    
    # 헤더 행 확인 (첫 번째 행이 헤더인지 판단)
    has_header = is_header_row(table_data[0]) if table_data else False
    
    # 표 너비 계산 (A4 페이지 너비에 맞춤)
    page_width = 8.27 * inch  # A4 너비
    margin = 0.5 * inch
    available_width = page_width - 2 * margin
    
    # 열 수에 따른 동적 너비 계산
    num_cols = len(table_data[0])
    col_width = available_width / num_cols
    
    # 표 생성
    table = Table(table_data, colWidths=[col_width] * num_cols)
    
    # 개선된 표 스타일
    if font_registered:
        table_style = TableStyle([
            # 격자 및 정렬
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            
            # 패딩
            ('TOPPADDING', (0, 0), (-1, -1), 6),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ('LEFTPADDING', (0, 0), (-1, -1), 4),
            ('RIGHTPADDING', (0, 0), (-1, -1), 4),
            
            # 텍스트 줄바꿈
            ('WORDWRAP', (0, 0), (-1, -1), True),
        ])
        
        # 헤더가 있으면 헤더 스타일 적용
        if has_header:
            table_style.add('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2E86AB'))
            table_style.add('TEXTCOLOR', (0, 0), (-1, 0), colors.white)
            table_style.add('FONTNAME', (0, 0), (-1, 0), 'KoreanFont')
            table_style.add('FONTSIZE', (0, 0), (-1, 0), 11)
            table_style.add('FONTWEIGHT', (0, 0), (-1, 0), 'bold')
            
            # 데이터 행 스타일
            table_style.add('FONTNAME', (0, 1), (-1, -1), 'KoreanFont')
            table_style.add('FONTSIZE', (0, 1), (-1, -1), 10)
            table_style.add('TEXTCOLOR', (0, 1), (-1, -1), colors.black)
            
            # 짝수 행 배경색 (가독성 향상)
            table_style.add('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#F8F9FA'))
        else:
            # 헤더가 없으면 모든 행에 동일한 스타일
            table_style.add('FONTNAME', (0, 0), (-1, -1), 'KoreanFont')
            table_style.add('FONTSIZE', (0, 0), (-1, -1), 10)
            table_style.add('TEXTCOLOR', (0, 0), (-1, -1), colors.black)
            table_style.add('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#F8F9FA'))
    else:
        # 기본 폰트 사용 (한국어가 깨질 수 있음)
        table_style = TableStyle([
            # 격자 및 정렬
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            
            # 패딩
            ('TOPPADDING', (0, 0), (-1, -1), 6),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ('LEFTPADDING', (0, 0), (-1, -1), 4),
            ('RIGHTPADDING', (0, 0), (-1, -1), 4),
            
            # 텍스트 줄바꿈
            ('WORDWRAP', (0, 0), (-1, -1), True),
        ])
        
        # 헤더가 있으면 헤더 스타일 적용
        if has_header:
            table_style.add('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2E86AB'))
            table_style.add('TEXTCOLOR', (0, 0), (-1, 0), colors.white)
            table_style.add('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold')
            table_style.add('FONTSIZE', (0, 0), (-1, 0), 11)
            
            # 데이터 행 스타일
            table_style.add('FONTNAME', (0, 1), (-1, -1), 'Helvetica')
            table_style.add('FONTSIZE', (0, 1), (-1, -1), 10)
            table_style.add('TEXTCOLOR', (0, 1), (-1, -1), colors.black)
            
            # 짝수 행 배경색 (가독성 향상)
            table_style.add('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#F8F9FA'))
        else:
            # 헤더가 없으면 모든 행에 동일한 스타일
            table_style.add('FONTNAME', (0, 0), (-1, -1), 'Helvetica')
            table_style.add('FONTSIZE', (0, 0), (-1, -1), 10)
            table_style.add('TEXTCOLOR', (0, 0), (-1, -1), colors.black)
            table_style.add('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#F8F9FA'))
    
    table.setStyle(table_style)
    return table

def is_header_row(row):
    """행이 헤더인지 확인"""
    if not row:
        return False
    
    # 헤더로 보이는 키워드들
    header_keywords = [
        '항목', '구분', '분류', '종류', '유형', '타입', '카테고리',
        '특성', '특징', '속성', '성질', '성격',
        '근거', '이유', '원인', '배경', '기반',
        '내용', '설명', '상세', '세부',
        '비율', '퍼센트', '수치', '값', '데이터',
        '날짜', '기간', '시기', '연도', '월', '일',
        '이름', '명칭', '제목', '표제',
        '번호', '순서', '순번', '인덱스'
    ]
    
    # 행의 모든 셀을 확인
    for cell in row:
        cell_lower = cell.lower().strip()
        for keyword in header_keywords:
            if keyword in cell_lower:
                return True
    
    # 셀 내용이 짧고 명확한 경우 (헤더일 가능성)
    short_cells = sum(1 for cell in row if len(cell.strip()) <= 10)
    if short_cells >= len(row) * 0.7:  # 70% 이상이 짧은 경우
        return True
    
    return False

def generate_pdf_report(content, user_inputs):
    """PDF 보고서 생성 - 표 처리 개선 및 폰트 문제 해결"""
    
    # 메모리 버퍼에 PDF 생성
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    
    # 스타일 설정
    styles = getSampleStyleSheet()
    
    # 한국어 폰트 등록
    font_registered = register_korean_font()
    
    # 커스텀 스타일 생성
    if font_registered:
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontName='KoreanFont',
            fontSize=16,
            spaceAfter=12,
            alignment=1  # 중앙 정렬
        )
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontName='KoreanFont',
            fontSize=14,
            spaceAfter=8
        )
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontName='KoreanFont',
            fontSize=10,
            spaceAfter=6
        )
    else:
        # 기본 폰트 사용 (한국어 지원 안됨)
        title_style = styles['Heading1']
        heading_style = styles['Heading2']
        normal_style = styles['Normal']
    
    # 내용을 문단으로 분할
    story = []
    
    # 제목 추가
    project_name = user_inputs.get('project_name', '프로젝트')
    title_text = f"{project_name} 분석 보고서"
    story.append(Paragraph(title_text, title_style))
    story.append(Spacer(1, 20))
    
    # 내용 파싱 및 추가
    paragraphs = content.split('\n\n')
    
    for para in paragraphs:
        para = para.strip()
        if not para:
            continue
            
        # 표 형식 처리
        if is_table_format(para):
            table_data, table_title = parse_table_from_text(para)
            if table_data and len(table_data) > 0:
                # 표 제목이 있으면 먼저 추가
                if table_title:
                    story.append(Paragraph(clean_text_for_pdf(table_title), heading_style))
                    story.append(Spacer(1, 6))
                
                # 표 생성 - 너비 제한으로 페이지 넘어김 방지
                table = create_table_with_improved_style(table_data, font_registered)
                
                if table:
                    story.append(table)
                    story.append(Spacer(1, 12))
                continue
        
        # 일반 텍스트 처리
        lines = para.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            # 제목 처리
            if line.startswith('# '):
                text = clean_text_for_pdf(line[2:].strip())
                story.append(Paragraph(text, title_style))
                story.append(Spacer(1, 12))
            elif line.startswith('## '):
                text = clean_text_for_pdf(line[3:].strip())
                story.append(Paragraph(text, heading_style))
                story.append(Spacer(1, 8))
            elif line.startswith('### '):
                text = clean_text_for_pdf(line[4:].strip())
                story.append(Paragraph(text, heading_style))
                story.append(Spacer(1, 6))
            elif line.startswith('---'):
                story.append(Spacer(1, 12))
            else:
                # 일반 텍스트 처리
                if line:
                    clean_line = clean_text_for_pdf(line)
                    if clean_line:
                        story.append(Paragraph(clean_line, normal_style))
    
    # PDF 생성
    try:
        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()
    except Exception as e:
        print(f"PDF 생성 오류: {e}")
        # 오류 발생 시 간단한 텍스트로 재시도
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4)
        simple_story = []
        
        # 간단한 텍스트로 변환
        simple_story.append(Paragraph(f"{project_name} 분석 보고서", title_style))
        simple_story.append(Spacer(1, 20))
        
        # 원본 텍스트를 그대로 사용 (HTML 태그 제거)
        clean_content = clean_text_for_pdf(content)
        paragraphs = clean_content.split('\n\n')
        
        for para in paragraphs:
            if para.strip():
                simple_story.append(Paragraph(para.strip(), normal_style))
                simple_story.append(Spacer(1, 6))
        
        doc.build(simple_story)
        buffer.seek(0)
        return buffer.getvalue()

def generate_word_report(content, user_inputs):
    """Word 문서 보고서 생성 - 표 처리 개선"""
    
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx 모듈이 설치되지 않았습니다. 'pip install python-docx'로 설치해주세요.")
    
    # Word 문서 생성
    doc = Document()
    
    # 제목 설정
    project_name = user_inputs.get('project_name', '프로젝트')
    title = doc.add_heading(f"{project_name} 분석 보고서", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 내용 파싱 및 추가
    paragraphs = content.split('\n\n')
    
    for para in paragraphs:
        para = para.strip()
        if not para:
            continue
            
        # 표 형식 처리
        if is_table_format(para):
            table_data, table_title = parse_table_from_text(para)
            if table_data and len(table_data) > 0:
                try:
                    # 표 제목이 있으면 먼저 추가
                    if table_title:
                        doc.add_heading(clean_text_for_pdf(table_title), level=3)
                    
                    # Word 표 생성
                    table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                    table.style = 'Table Grid'
                    
                    # 표 스타일 개선
                    table.allow_autofit = True
                    
                    # 데이터 채우기
                    for i, row in enumerate(table_data):
                        for j, cell in enumerate(row):
                            if i < len(table.rows) and j < len(table.rows[i].cells):
                                cell_text = clean_text_for_pdf(cell)
                                table.rows[i].cells[j].text = cell_text
                                
                                # 헤더 행 스타일링
                                if i == 0:
                                    cell_obj = table.rows[i].cells[j]
                                    for paragraph in cell_obj.paragraphs:
                                        for run in paragraph.runs:
                                            run.bold = True
                                            run.font.color.rgb = None  # 기본 색상
                
                except Exception as e:
                    print(f"표 생성 오류: {e}")
                    # 표 생성 실패 시 일반 텍스트로 변환
                    if table_title:
                        doc.add_paragraph(f"[표 제목: {table_title}]")
                    doc.add_paragraph(f"[표 데이터: {para[:100]}...]")
                
                doc.add_paragraph()  # 표 후 빈 줄
                continue
        
        # 일반 텍스트 처리
        lines = para.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            # 제목 처리
            if line.startswith('# '):
                text = clean_text_for_pdf(line[2:].strip())
                doc.add_heading(text, level=1)
            elif line.startswith('## '):
                text = clean_text_for_pdf(line[3:].strip())
                doc.add_heading(text, level=2)
            elif line.startswith('### '):
                text = clean_text_for_pdf(line[4:].strip())
                doc.add_heading(text, level=3)
            elif line.startswith('---'):
                doc.add_paragraph()  # 빈 줄 추가
            else:
                # 일반 텍스트 처리
                if line:
                    clean_line = clean_text_for_pdf(line)
                    if clean_line:
                        doc.add_paragraph(clean_line)
    
    # 메모리 버퍼에 저장
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def generate_report_content(report_type, include_charts, include_recommendations, include_appendix):
    """보고서 내용 생성 - 보고서 유형별 차이점 적용"""
    from user_state import get_user_inputs
    user_inputs = get_user_inputs()
    
    # 기본 정보
    report_content = f"""
# {user_inputs.get('project_name', '프로젝트')} 분석 보고서
**보고서 유형**: {report_type}

## 📋 프로젝트 기본 정보
- **프로젝트명**: {user_inputs.get('project_name', 'N/A')}
- **건축주**: {user_inputs.get('owner', 'N/A')}
- **대지위치**: {user_inputs.get('site_location', 'N/A')}
- **대지면적**: {user_inputs.get('site_area', 'N/A')}
- **건물용도**: {user_inputs.get('building_type', 'N/A')}
- **프로젝트 목표**: {user_inputs.get('project_goal', 'N/A')}

"""
    
    # 분석 결과 추가
    import streamlit as st
    if st.session_state.get('cot_history'):
        if report_type == "전체 분석 보고서":
            # 전체 분석 보고서: 모든 상세 내용 포함
            report_content += "## 전체 분석 결과\n"
            for i, history in enumerate(st.session_state.cot_history, 1):
                report_content += f"""
### {i}. {history['step']}

**요약**: {history.get('summary', '')}

**인사이트**: {history.get('insight', '')}

**상세 분석**:
{history.get('result', '')}

---
"""
        
        elif report_type == "요약 보고서":
            # 요약 보고서: 핵심 요약과 인사이트만
            for i, history in enumerate(st.session_state.cot_history, 1):
                report_content += f"""
### {i}. {history['step']}

**핵심 요약**: {history.get('summary', '')}

**주요 인사이트**: {history.get('insight', '')}

---
"""
        
        elif report_type == "전문가 보고서":
            # 전문가 보고서: 기술적 분석과 전문적 권장사항
            report_content += "## 전문가 분석 결과\n"
            for i, history in enumerate(st.session_state.cot_history, 1):
                report_content += f"""
### {i}. {history['step']}

**분석 요약**: {history.get('summary', '')}

**전문가 인사이트**: {history.get('insight', '')}

**기술적 분석**:
{history.get('result', '')[:500]}...

---
"""
        
        elif report_type == "클라이언트 보고서":
            # 클라이언트 보고서: 비즈니스 관점의 핵심 내용
            report_content += "## 💼 비즈니스 분석 결과\n"
            for i, history in enumerate(st.session_state.cot_history, 1):
                report_content += f"""
### {i}. {history['step']}

**비즈니스 요약**: {history.get('summary', '')}

**핵심 가치**: {history.get('insight', '')}

**실행 가능한 제안**:
{history.get('result', '')[:300]}...

---
"""
    
    # 추가 섹션들 (보고서 유형별로 다르게 적용)
    if include_charts:
        if report_type == "전체 분석 보고서":
            report_content += """
## 상세 차트 및 다이어그램
(모든 차트 및 다이어그램이 포함됩니다)
"""
        elif report_type == "요약 보고서":
            report_content += """
## 핵심 차트
(주요 차트만 포함됩니다)
"""
        elif report_type == "전문가 보고서":
            report_content += """
## 전문가 차트 및 분석 다이어그램
(기술적 분석을 위한 상세 차트가 포함됩니다)
"""
        elif report_type == "클라이언트 보고서":
            report_content += """
## 💼 비즈니스 차트
(비즈니스 의사결정을 위한 핵심 차트가 포함됩니다)
"""
    
    if include_recommendations:
        if report_type == "전체 분석 보고서":
            report_content += """
## 💡 종합 권장사항
분석 결과를 바탕으로 한 상세한 권장사항이 포함됩니다.
"""
        elif report_type == "요약 보고서":
            report_content += """
## 💡 핵심 권장사항
가장 중요한 권장사항만 포함됩니다.
"""
        elif report_type == "전문가 보고서":
            report_content += """
## 💡 전문가 권장사항
기술적 관점에서의 전문적 권장사항이 포함됩니다.
"""
        elif report_type == "클라이언트 보고서":
            report_content += """
## 💡 비즈니스 권장사항
비즈니스 관점에서의 실행 가능한 권장사항이 포함됩니다.
"""
    
    if include_appendix:
        if report_type == "전체 분석 보고서":
            report_content += """
## 📋 상세 부록
모든 추가 자료 및 참고문헌이 포함됩니다.
"""
        elif report_type == "요약 보고서":
            report_content += """
## 📋 핵심 부록
주요 참고자료만 포함됩니다.
"""
        elif report_type == "전문가 보고서":
            report_content += """
## 📋 전문가 부록
기술적 참고자료 및 전문 문헌이 포함됩니다.
"""
        elif report_type == "클라이언트 보고서":
            report_content += """
## 📋 비즈니스 부록
비즈니스 관련 참고자료가 포함됩니다.
"""
    
    return report_content